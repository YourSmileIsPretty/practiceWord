import docx 
import pymorphy2
import datetime
import enchant
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
import re 

doc = docx.Document('title_document.docx')

#Creating an instance of a class for morphological analysis of words
morph = pymorphy2.MorphAnalyzer()

for para in doc.paragraphs:
    words = para.text.split()
    new_words = []
    for word in words:
        # Проверием, является ли входное значение именем собственным 
        parsed_word = morph.parse(word)[0]
        if 'Name' in parsed_word.tag:
            new_words.append(word)
        else:
            # Заменяем буквы "Ё" на "Е" 
            new_word = word.replace('Ё', 'е').replace('Ё', 'е') # Вызываем ещё раз так как может быть случай, когда буква "Ё" может быть заглавной
            new_words.append(new_word)
    # Сохраняем изменённый текст в новый документ Word 
    new_text = ' '.join(new_words)
    new_para = doc.add_paragraph(new_text)
    print("Все буквы Ё были исправлены")

#----------------------------------------------------------------------------------------------------
# Узнаём текущий год
current_year = datetime.datetime.now().year
current_number = '1'

# Открываем документ
doc = docx.Document('title_document.docx')

# Получаем первый колонтитул и его текст
header = doc.sections[0].header
header_text = header.text

# Находим год и номер в тексте колонтитула
if '2023' in header_text or '1' in header_text:
    # Заменяем год и номер в тексте колонтитула на текущие значения
    new_header_text = header_text.replace('2023', str(current_year)).replace('1', current_number)

    # Заменяем текст колонтитула на новый текст
    header.text = new_header_text

    # Сохраняем изменения
    doc.save('title_document.docx')

    print('Год и номер заменены')
else:
    print('Год и номер уже актуальны')

#----------------------------------------------------------------------------------------------------
# Задаём актуальные выражения реквезитов
current_vid = 'Договор'
current_date = '25.04.2023'
current_number = '1'
current_name = 'Название документа'

# Открываем документ
doc = docx.Document('title_document.docx')

# Перебираем все параграфы
for para in doc.paragraphs:
    # Проверяем, что параграф выровнен по центру
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        # Получаем текст параграфа
        text = para.text.strip()
        # Разбиваем текст по символу "—"
        parts = text.split('—')
        # Проверяем, что количество частей равно 4
        if len(parts) == 4:
            # Получаем реквизиты
            kind, date, number, name = [part.strip() for part in parts]
            # Проверяем корректность реквизитов
            if kind and date and number and name:
                # Заменяем название документа на актуальное
                if name == 'Название документа':
                    name = 'Актуальное название документа'
                    para.text = f'{current_vid} — {current_date} — {current_number} — {current_name}'
                    # Выводим информацию об изменении
                    print(f'Реквизиты документа в параграфе "{text}" были изменены на "{para.text}"')
                    
# Сохраняем изменения в документе
doc.save('title_document.docx')
#----------------------------------------------------------------------------------------------------
#Отсутствуют множественные пробелы между словами (проверяйте, включая отображение невидимых знаков)
# Открываем документ Word
doc = docx.Document('title_document.docx')

# Проходимся по каждому абзацу в документе
for paragraph in doc.paragraphs:
    # Получаем текст абзаца и заменяем все невидимые символы на обычные пробелы
    text = paragraph.text.replace('\t', ' ').replace('\n', ' ').replace('\r', ' ')
    while '  ' in text:
        text = text.replace('  ', ' ')

    # Проверяем наличие множественных пробелов между словами
    if '  ' in text:
        text = re.sub(r'\s+', ' ', text)

    
    # Добавляем разрыв после абзаца
    paragraph.add_run().add_break(WD_BREAK.LINE)
    
# Сохраняем изменения в документе
doc.save('title_document.docx')
#----------------------------------------------------------------------------------------------------
# Проверить чтобы интервал шрифта во всём файле были одинаковые 
# 1 вариант решения создание своего собственного стиля с самого начала 
    # doc = docx.Document('title_document.docx')
    # style = doc.styles.add_style('CustomStyle', doc.styles['Normal'])
    # style.paragraph_format.line_spacing = 1.5 # задаем интервал между строками
    # for paragraph in doc.paragraphs:
    #     paragraph.style = style # применяем стиль ко всем параграфам в документе
    # doc.save('title_document.docx')

# 2 вариант решения проверить различается ли интервал во всём файле и применить интервал, который первый попадается в выборку 
# Открываем документ и получаем все его параграфы
doc = docx.Document('title_document.docx')
paragraphs = doc.paragraphs

# Получаем текущий интервал документа
current_spacing = paragraphs[0].paragraph_format.line_spacing
for p in paragraphs:
    if p.paragraph_format.line_spacing != current_spacing:
        # Если интервал отличается от текущего, заменяем его
        p.paragraph_format.line_spacing = docx.enum.text.WD_LINE_SPACING.ONE_POINT_FIVE

# Сохраняем изменения в документе
doc.save('title_document.docx')